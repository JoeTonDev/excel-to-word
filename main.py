import pandas as pd
from docx import Document

# Load excel file into a pandas dataframe
file_path = "property_data.xlsx"
df = pd.read_excel(file_path)

# Create a new document
document = Document()

# Iterate through each row in the dataframe
for i, row in df.iterrows():
  # Get the property data from the current row
  property_id = row["Property ID"]
  address = row["Address"]
  owner = row["Owner"]
  tenant = row["Tenant"]
  
# Add new agreement section to the document
document.add_heading(f"Agreement for Property ID: {property_id}", lelvel=1)
document.add_paragraph(f"Address: {address}")
document.add_paragraph(f"This agreement is made between the property owner: {owner}, and the tenant: {tenant}.")

# Save the document
document.save("property_agreements.docx")